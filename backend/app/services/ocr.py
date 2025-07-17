from langchain.schema import Document
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import docx
import os
from io import BytesIO

def process_document(file_obj, filename="uploaded_file"):
    name = filename
    ext = os.path.splitext(name)[1].lower()

    text_chunks = []

    if ext in [".jpg", ".jpeg", ".png"]:
        image = Image.open(file_obj)
        text = pytesseract.image_to_string(image, lang="eng")
        text_chunks = [Document(page_content=text, metadata={"source": name, "page": 1})]

    elif ext == ".pdf":
        content = file_obj.read()
        pdf = fitz.open(stream=content, filetype="pdf")
        for i, page in enumerate(pdf):
            text = page.get_text().strip()
            if text:
                text_chunks.append(Document(page_content=text, metadata={"source": name, "page": i + 1}))
        file_obj.seek(0)  # Reset pointer if needed later

    elif ext == ".docx":
        # If file_obj is a SpooledTemporaryFile or bytes, ensure docx can handle it
        if hasattr(file_obj, 'read'):
            docx_obj = docx.Document(file_obj)
        else:
            docx_obj = docx.Document(BytesIO(file_obj))
        full_text = "\n".join([para.text for para in docx_obj.paragraphs if para.text.strip()])
        text_chunks = [Document(page_content=full_text, metadata={"source": name, "page": 1})]

    elif ext == ".txt":
        text = file_obj.read().decode("utf-8")
        text_chunks = [Document(page_content=text, metadata={"source": name, "page": 1})]

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return text_chunks